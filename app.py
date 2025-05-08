import nltk
from nltk.corpus import stopwords
from nltk.stem.snowball import SnowballStemmer
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.cluster import KMeans
from sklearn.decomposition import LatentDirichletAllocation
import numpy as np
import pandas as pd
import matplotlib
import base64
from io import BytesIO
import os
from flask import Flask, render_template, request
from docx import Document


matplotlib.use("Agg")
import matplotlib.pyplot as plt
import seaborn as sns

os.environ["LOKY_MAX_CPU_COUNT"] = "1"

nltk.download("punkt")
nltk.download("punkt_tab")
nltk.download("stopwords")

app = Flask(__name__)

stop_words = set(stopwords.words("russian"))
stemmer = SnowballStemmer("russian")


def extract_text_from_file(file):
    filename = file.filename.lower()

    if filename.endswith(".txt"):
        try:
            return file.read().decode("utf-8")
        except UnicodeDecodeError:
            file.seek(0)
            return file.read().decode("cp1251")
    elif filename.endswith(".docx"):
        doc = Document(file)
        return "\n".join([para.text for para in doc.paragraphs])
    elif filename.endswith(".doc"):
        try:
            doc = Document(file)
            return "\n".join([para.text for para in doc.paragraphs])
        except:
            return (
                "Не удалось прочитать .doc файл. Пожалуйста, используйте .docx или .txt"
            )
    else:
        raise ValueError("Неподдерживаемый формат файла")


def preprocess_text(text):
    tokens = nltk.word_tokenize(text, language="russian")
    tokens = [token.lower() for token in tokens if token.isalnum()]
    tokens = [stemmer.stem(token) for token in tokens if token not in stop_words]
    return " ".join(tokens)


def fig_to_base64(fig):
    buf = BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight", dpi=120)
    plt.close(fig)
    return base64.b64encode(buf.getvalue()).decode("utf-8")


def compute_tfidf(docs):
    vectorizer = TfidfVectorizer()
    X = vectorizer.fit_transform(docs)
    return X, vectorizer.get_feature_names_out()


def find_optimal_clusters(tfidf_matrix):
    num_samples = tfidf_matrix.shape[0]
    max_clusters = min(num_samples, 10)
    if max_clusters < 2:
        return 2, None

    wcss = []
    for i in range(1, max_clusters + 1):
        kmeans = KMeans(
            n_clusters=i, init="k-means++", max_iter=300, n_init=10, random_state=42
        )
        kmeans.fit(tfidf_matrix)
        wcss.append(kmeans.inertia_)

    if len(wcss) < 3:
        return 2, None

    elbow_diff = np.diff(wcss)
    elbow_diff_ratio = elbow_diff[1:] / elbow_diff[:-1]

    if len(elbow_diff_ratio) == 0:
        return 2, None

    optimal_clusters = np.argmin(elbow_diff_ratio) + 2
    optimal_clusters = max(2, min(optimal_clusters, max_clusters))

    fig = plt.figure(figsize=(10, 7))
    plt.plot(range(1, max_clusters + 1), wcss, marker="o")
    plt.axvline(x=optimal_clusters, color="r", linestyle="--")
    plt.title("Метод локтя")
    plt.xlabel("Количество кластеров")
    plt.ylabel("WCSS (within-cluster sums of squares)")

    return optimal_clusters, fig


def cluster_texts(X, n_clusters=4):
    kmeans = KMeans(
        n_clusters=n_clusters,
        init="k-means++",
        max_iter=300,
        n_init=10,
        random_state=42,
    )
    clusters = kmeans.fit_predict(X)
    return clusters


def reduce_dimensions(X, n_components=4):
    lda = LatentDirichletAllocation(n_components=n_components, random_state=42)
    X_reduced = lda.fit_transform(X)
    return X_reduced


def visualize_clusters(reduced_matrix, clusters):
    fig = plt.figure(figsize=(10, 7))
    scatter = plt.scatter(
        reduced_matrix[:, 0], reduced_matrix[:, 1], c=clusters, cmap="viridis", s=100
    )
    plt.colorbar(scatter)
    plt.xlabel("Главная компонента 1")
    plt.ylabel("Главная компонента 2")
    plt.title("Визуализация кластеров документов")
    return fig


def additional_visualizations(tfidf_matrix, clusters, feature_names):
    df_tfidf = pd.DataFrame(tfidf_matrix.toarray(), columns=feature_names)
    df_tfidf["cluster"] = clusters

    fig = plt.figure(figsize=(10, 6))
    sns.countplot(
        x="cluster", data=df_tfidf, hue="cluster", palette="viridis", legend=False
    )
    plt.title("Распределение документов по кластерам")
    plt.xlabel("Кластеры")
    plt.ylabel("Количество документов")

    for p in plt.gca().patches:
        plt.gca().annotate(
            f"{int(p.get_height())}",
            (p.get_x() + p.get_width() / 2.0, p.get_height()),
            ha="center",
            va="center",
            xytext=(0, 5),
            textcoords="offset points",
        )
    return fig


def analyze_clusters(processed_docs, clusters):
    cluster_analysis = {}
    for cluster_num in set(clusters):
        cluster_docs = [
            processed_docs[i] for i, c in enumerate(clusters) if c == cluster_num
        ]
        all_words = " ".join(cluster_docs).split()
        word_counts = pd.Series(all_words).value_counts().head(10)
        cluster_analysis[cluster_num] = word_counts.to_dict()
    return cluster_analysis


@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        if "documents[]" not in request.files:
            return render_template(
                "index.html",
                error="Пожалуйста, загрузите хотя бы один файл",
                documents=[],
            )

        files = request.files.getlist("documents[]")
        documents = []
        filenames = []

        for file in files:
            if file.filename == "":
                continue

            filename = file.filename.lower()
            if not (
                filename.endswith(".txt")
                or filename.endswith(".docx")
                or filename.endswith(".doc")
            ):
                return render_template(
                    "index.html",
                    error=f"Неподдерживаемый формат файла: {file.filename}. Разрешены только .txt, .docx, .doc",
                    documents=[],
                )

            try:
                text = extract_text_from_file(file)
                if text.strip():
                    documents.append(text)
                    filenames.append(file.filename)
            except Exception as e:
                return render_template(
                    "index.html",
                    error=f"Ошибка при чтении файла {file.filename}: {str(e)}",
                    documents=[],
                )

        if not documents:
            return render_template(
                "index.html",
                error="Пожалуйста, загрузите хотя бы один файл с текстом",
                documents=[],
            )

        try:
            processed_docs = [preprocess_text(doc) for doc in documents]
            tfidf_matrix, feature_names = compute_tfidf(processed_docs)

            optimal_clusters, elbow_plot_fig = find_optimal_clusters(tfidf_matrix)
            clusters = cluster_texts(tfidf_matrix, n_clusters=optimal_clusters)
            reduced_matrix = reduce_dimensions(
                tfidf_matrix, n_components=optimal_clusters
            )

            elbow_plot = fig_to_base64(elbow_plot_fig) if elbow_plot_fig else None
            cluster_plot = fig_to_base64(visualize_clusters(reduced_matrix, clusters))
            distribution_plot = fig_to_base64(
                additional_visualizations(tfidf_matrix, clusters, feature_names)
            )

            cluster_words = analyze_clusters(processed_docs, clusters)

            unique, counts = np.unique(clusters, return_counts=True)
            cluster_distribution = {
                f"Кластер {int(cluster)}": int(count)
                for cluster, count in zip(unique, counts)
            }

            doc_cluster_mapping = [
                {
                    "document": filenames[i],
                    "text": documents[i],
                    "cluster": int(cluster),
                }
                for i, cluster in enumerate(clusters)
            ]

            return render_template(
                "index.html",
                elbow_plot=elbow_plot,
                cluster_plot=cluster_plot,
                distribution_plot=distribution_plot,
                cluster_words=cluster_words,
                doc_cluster_mapping=doc_cluster_mapping,
                cluster_distribution=cluster_distribution,
                feature_names=feature_names.tolist(),
                processed_docs=processed_docs,
                optimal_clusters=optimal_clusters,
                documents=[],
                show_results=True,
            )
        except Exception as e:
            return render_template(
                "index.html",
                error=f"Произошла ошибка при обработке документов: {str(e)}",
                documents=[],
            )

    return render_template("index.html", documents=[])


if __name__ == "__main__":
    app.run(debug=True)
